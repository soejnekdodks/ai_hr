import io
from docx import Document
import docx2txt
from striprtf.striprtf import rtf_to_text
import fitz


def txt_to_text(txt_bytes: bytes) -> str:
    try:
        # Пытаемся декодировать с разными кодировками
        encodings = ['utf-8', 'windows-1251', 'cp866', 'iso-8859-1', 'utf-16']
        
        for encoding in encodings:
            try:
                text = txt_bytes.decode(encoding)
                return text
            except UnicodeDecodeError:
                continue
        
        # Если ни одна кодировка не подошла, пробуем с игнорированием ошибок
        return txt_bytes.decode('utf-8', errors='ignore')
        
    except Exception as e:
        print(f"Ошибка при чтении TXT: {e}")
        return ""


def pdf_to_text(pdf_bytes: bytes) -> str:
    try:
        text = ""
        with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
            for page in doc:
                text += page.get_text() + "\n"
        return text
    except Exception as e:
        print(f"Ошибка при чтении PDF: {e}")
        return ""


def docx_to_text(docx_bytes: bytes) -> str:
    try:
        docx_stream = io.BytesIO(docx_bytes)
        doc = Document(docx_stream)
        
        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
        
        return "\n".join(text)
    except Exception as e:
        print(f"Ошибка при чтении DOCX: {e}")
        return ""


def rtf_to_text_bytes(rtf_bytes: bytes) -> str:
    try:
        # Декодируем байты в строку (RTF обычно в кодировке Windows-1252 или cp866)
        try:
            rtf_content = rtf_bytes.decode('windows-1252')
        except UnicodeDecodeError:
            try:
                rtf_content = rtf_bytes.decode('cp866')
            except UnicodeDecodeError:
                rtf_content = rtf_bytes.decode('utf-8', errors='ignore')
        
        text = rtf_to_text(rtf_content)
        return text
    except Exception as e:
        print(f"Ошибка при чтении RTF: {e}")
        return ""


def document_to_text(file_bytes: bytes, file_extension: str) -> str:
    file_extension = file_extension.lower().lstrip('.')
    
    try:
        if file_extension == 'pdf':
            return pdf_to_text(file_bytes)
        
        elif file_extension in ['docx', 'doc']:
            return docx_to_text(file_bytes)
        
        elif file_extension == 'rtf':
            return rtf_to_text_bytes(file_bytes)
        
        elif file_extension == 'txt':
            return txt_to_text(file_bytes)
        
        else:
            print(f"Неподдерживаемый формат: {file_extension}")
            return ""
            
    except Exception as e:
        print(f"Ошибка при обработке {file_extension}: {e}")
        return ""